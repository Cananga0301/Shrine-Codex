"""DOCX parser for Vietnamese legal documents."""

from __future__ import annotations

import logging
from pathlib import Path
from docx import Document as DocxDocument

log = logging.getLogger(__name__)


def read_docx(file_path: str | Path) -> str:
    """Extract full text from a DOCX file via python-docx."""
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if file_path.suffix.lower() == ".doc":
        raise RuntimeError("Parsing .doc files is not supported in the Docker build. Convert to .docx first.")

    return _parse_docx(file_path)


def parse_docx(file_path: str | Path) -> str:
    """Backward-compatible alias for legacy imports."""
    return read_docx(file_path)


def _parse_docx(file_path: Path) -> str:
    """Parse .docx via python-docx."""
    doc = DocxDocument(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    log.info("Parsed DOCX '%s': %d characters.", file_path.name, len(full_text))
    return full_text


